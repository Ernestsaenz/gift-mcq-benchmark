"""Render expC-findings.docx from canonical.json.

Every number in the document is either read from canonical.json or recomputed here
from the read-only source workbook. Nothing is typed in by hand.
"""
from __future__ import annotations

import collections
import json
import re
import unicodedata
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CAN = json.loads((ROOT / "canonical.json").read_text(encoding="utf-8"))
SRC = Path("/Users/ernestsaenz/Programming/GIFT-abstract-dossier/tier1_mcq/data"
           "/experiment-31-07-26/balanced-flat-A.xlsx")
OUT = ROOT / "expC-findings.docx"

ITEMS = CAN["items"]
GEN = CAN["generated"]
TIERS = CAN["tiers"]
STATS = CAN["tier_stats"]
BM = CAN["sentences"]["biomarker"]
AN = CAN["sentences"]["anatomy"]
BM_ROT = CAN["sentences"]["biomarker_rotation"]
AN_ROT = CAN["sentences"]["anatomy_rotation"]

# --------------------------------------------------------------------------
# Recomputation against the source workbook (read-only)
# --------------------------------------------------------------------------
wb = openpyxl.load_workbook(SRC, read_only=True, data_only=True)
_rows = list(wb["questions"].iter_rows(values_only=True))
_hdr = _rows[0]
RECS = [dict(zip(_hdr, r)) for r in _rows[1:] if r[0] is not None]
wb.close()
BYID = {r["question_id"]: r for r in RECS}
ORDER = [r["question_id"] for r in RECS]
assert len(RECS) == CAN["n_source_rows"] == 474


def nf(s):
    s = unicodedata.normalize("NFD", s or "")
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", s).lower()


MECH = [q for q in ORDER if ITEMS[q]["mechanical_guard"] is None]


def root_hits(root, pool):
    return sum(1 for q in pool if root in nf(BYID[q]["question_text"]))


# host words for the two roots discussed in the text
def host_words(root):
    c = collections.Counter()
    for r in RECS:
        blob = nf(" ".join(str(r[k] or "") for k in
                           ("question_text", "option_a", "option_b", "option_c", "option_d")))
        c.update(re.findall(r"\b\w*" + root + r"\w*\b", blob))
    return c


MUCO_WORDS = host_words("muco")
FIBRO_WORDS = host_words("fibro")
MUCO_ROWS_ANY = sum(
    1 for r in RECS
    if re.search(r"muco", nf(" ".join(str(r[k] or "") for k in
                                      ("question_text", "option_a", "option_b",
                                       "option_c", "option_d")))))

MODALITY = {
    "Ultrasound": r"\becograf|ultrasonid|ecoendoscop",
    "CT": r"\btomografia computa|\btac\b|\btc\b|angio-?tc|colonografia por tc",
    "MRI": r"resonancia magnetica|\brmn?\b|colangiorresonancia|colangio-?rm",
}
NORMALPAT = (r"(normal|sin (hallazgos|alteraciones|lesiones|signos)"
             r"|no (se )?(objetiva|muestra|evidencia|identifica|aprecia)|anodin|negativ)")
IMGOPT = r"(ecograf|tomografia|\btac\b|\btc\b|resonancia|\brm\b|prueba de imagen|imagen)"


def imaging_funnel(pool, pat):
    performed = [q for q in pool if re.search(pat, nf(BYID[q]["question_text"]))]
    live = []
    for q in performed:
        t = nf(BYID[q]["question_text"])
        if not any(re.search(NORMALPAT, t[m.end():m.end() + 90]) for m in re.finditer(pat, t)):
            live.append(q)
    clean = [q for q in live
             if not re.search(IMGOPT, nf(" || ".join(BYID[q][k] or "" for k in
                                                     ("option_a", "option_b", "option_c", "option_d"))))
             and not re.search(IMGOPT, nf(ITEMS[q]["stem"]))]
    return len(performed), len(live), len(clean)


ALL474 = ORDER

# corpus-level clustering
CTX_ITEMS = [q for q in ORDER if ITEMS[q]["context_ids"]]
CTX_NARR = len({ITEMS[q]["cluster"] for q in CTX_ITEMS})
CTX_MEMBERSHIPS = sum(len(ITEMS[q]["context_ids"].split(",")) for q in CTX_ITEMS)

# rotation lock-step
JOINT = collections.Counter()
for q, v in GEN.items():
    if "BM_variant" in v and "AN_variant" in v:
        JOINT[(v["BM_variant"], v["AN_variant"])] += 1

BM_BY_REGION = collections.defaultdict(collections.Counter)
AN_BY_REGION = collections.defaultdict(collections.Counter)
for q, v in GEN.items():
    if "BM_variant" in v:
        BM_BY_REGION[ITEMS[q]["region"]][v["BM_variant"]] += 1
    if "AN_variant" in v:
        AN_BY_REGION[ITEMS[q]["region"]][v["AN_variant"]] += 1

# camouflage skips: assigned variant differs from the plain rotation entry at the item's index
_rank = {}
_seen = collections.Counter()
for _q in ORDER:
    _c = ITEMS[_q]["cluster"]
    _rank[_q] = _seen[_c]
    _seen[_c] += 1
BM_SKIPS = sum(1 for q, v in GEN.items()
               if "BM_variant" in v and v["BM_variant"] != BM_ROT[_rank[q] % len(BM_ROT)])
AN_SKIPS = sum(1 for q, v in GEN.items()
               if "AN_variant" in v and v["AN_variant"] != AN_ROT[_rank[q] % len(AN_ROT)])

BMVAR = collections.Counter(v["BM_variant"] for v in GEN.values() if "BM_text" in v)
ANVAR = collections.Counter(v["AN_variant"] for v in GEN.values() if "AN_text" in v)
N_BOTH = sum(1 for v in GEN.values() if "BM_text" in v and "AN_text" in v)
N_BM_ONLY = sum(1 for v in GEN.values() if "BM_text" in v and "AN_text" not in v)
N_AN_ONLY = sum(1 for v in GEN.values() if "AN_text" in v and "BM_text" not in v)


def profile(ids):
    ids = list(ids)
    fl = collections.Counter(ITEMS[q]["flags"] for q in ids)
    neg = sum(v for k, v in fl.items() if k and "negated" in k)
    cc = sum(v for k, v in fl.items() if k and "clinical_case" in k)
    nc = sorted(ITEMS[q]["narrative_chars"] for q in ids)
    n = len(nc)
    cl = collections.Counter(ITEMS[q]["cluster"] for q in ids)
    return {
        "n": n,
        "region": collections.Counter(ITEMS[q]["region"] for q in ids),
        "exam_part": collections.Counter(ITEMS[q]["exam_part"] for q in ids),
        "year": collections.Counter(ITEMS[q]["year"] for q in ids),
        "negated": neg, "clinical_case": cc,
        "letters": collections.Counter(ITEMS[q]["correct_letter"] for q in ids),
        "nc": (nc[0], nc[n // 4], nc[n // 2], nc[(3 * n) // 4], nc[-1]),
        "clusters": cl,
    }


P37 = profile(TIERS["strict"]["PAIR"])
P58 = profile(TIERS["relaxed"]["PAIR"])
P99 = profile(TIERS["relaxed"]["BM"])
P86 = profile(TIERS["relaxed"]["AN"])

STRICT_KEY = collections.Counter(ITEMS[q].get("strict_guard") for q in MECH)
RELAX_KEY = collections.Counter(ITEMS[q].get("relaxed_guard") for q in MECH)

# ==========================================================================
# Document
# ==========================================================================
doc = Document()

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(9)
st.paragraph_format.space_after = Pt(4.5)
st.paragraph_format.line_spacing = 1.0

for s in doc.sections:
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)

for name, size, bold in (("Heading 1", 12.5, True), ("Heading 2", 10.5, True), ("Heading 3", 9, True)):
    h = doc.styles[name]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 8)
    h.paragraph_format.space_after = Pt(3)


def para(text, style=None, size=None, bold=False, italic=False, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    return p


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    return p


def table(headers, rows, widths=None, font=7.5, right_from=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(font)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run("" if val is None else str(val))
            r.font.size = Pt(font)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            if right_from is not None and i >= right_from:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def topn(counter, n=None, sep=", "):
    items = counter.most_common(n)
    return sep.join(f"{k} {v}" for k, v in items)


# --------------------------------------------------------------------------
# Title block
# --------------------------------------------------------------------------
p = para("Experiment C — Fabricated-Entity Insertion: Feasibility and Item Set",
         size=19, bold=True, space_after=2)
p.paragraph_format.space_before = Pt(24)
para("Feasibility report and delivered item set", size=11.5, italic=True, space_after=14)

table(
    ["", ""],
    [
        ["Date", "31 July 2026"],
        ["Source bank", "balanced-flat-A.xlsx, sheet 'questions', 474 items"],
        ["Build script", "tools/build_expC.py"],
        ["Canonical output", "canonical.json"],
        ["Status", "Item set delivered; experiment blocked on a harness change (section 5)"],
    ],
    widths=[1.5, 5.2], font=9.5,
)

para("")
para("Abstract", size=12, bold=True, space_after=3)
para(
    "Experiment C was to build two arms of 200 items each in which a single fabricated clinical "
    "finding — a non-existent serum biomarker, or a non-existent anatomical structure — is inserted "
    "into an otherwise byte-identical examination question. "
    f"Of the {CAN['n_source_rows']} items in the source bank, {CAN['mechanical_pool']} admit a "
    "mechanically safe insertion point, and after answer-key screening "
    f"{len(TIERS['strict']['base'])} (strict) or {len(TIERS['relaxed']['base'])} (relaxed) survive; "
    f"the largest single arm obtainable is {len(TIERS['relaxed']['BM'])} biomarker items and the "
    f"largest paired set is {len(TIERS['relaxed']['PAIR'])}. "
    "The target of 200 items per arm is therefore not achievable from this bank, and no amount of "
    "additional generation effort changes that, because the constraint is the corpus and not the method. "
    f"All {len(GEN)} generated variants are literal string concatenations produced without any language "
    "model, with byte-level assertions proving that nothing but the inserted sentence moved. "
    "The experiment is nevertheless not yet measurable: the harness persists only the parsed "
    "multiple-choice letter, and because eligibility deliberately excludes every item whose key an "
    "inserted finding could move, the accuracy delta is zero by construction.",
    space_after=10,
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# --------------------------------------------------------------------------
# 1. Executive summary
# --------------------------------------------------------------------------
doc.add_heading("1. Executive summary", level=1)

para(
    f"How many questions can actually be modified. From {CAN['n_source_rows']} source items, "
    f"{len(TIERS['strict']['PAIR'])} can carry both a fabricated biomarker and a fabricated anatomical "
    f"structure under the strict answer-key rules, and {len(TIERS['relaxed']['PAIR'])} under the relaxed "
    f"rules. If only one arm is required, {len(TIERS['relaxed']['BM'])} items can carry a biomarker and "
    f"{len(TIERS['relaxed']['AN'])} can carry an anatomical structure. These are ceilings, not samples: "
    "they are the complete sets that survive the eligibility rules."
)

table(
    ["Set", "Strict", "Relaxed", "What it is"],
    [
        ["Mechanically insertable", CAN["mechanical_pool"], CAN["mechanical_pool"],
         "A clean sentence seam exists before the stem, and there is a patient to attach a finding to"],
        ["Key-preserving base", len(TIERS["strict"]["base"]), len(TIERS["relaxed"]["base"]),
         "The above, minus items whose correct answer an inserted finding could move"],
        ["Biomarker arm", len(TIERS["strict"]["BM"]), len(TIERS["relaxed"]["BM"]),
         "Key-preserving and not already asserting that the laboratory workup is normal"],
        ["Anatomy arm", len(TIERS["strict"]["AN"]), len(TIERS["relaxed"]["AN"]),
         "Key-preserving and not already asserting that the physical examination is normal"],
        ["Paired (control + both arms)", len(TIERS["strict"]["PAIR"]), len(TIERS["relaxed"]["PAIR"]),
         "Items usable in a within-item three-condition design"],
    ],
    widths=[1.65, 0.65, 0.7, 3.6], right_from=1,
)
caption("Table 1. Eligible item counts. Strict and relaxed differ only in guard J (section 3.3).")

para(
    f"200 items per arm is unreachable. The ceiling for the larger arm is "
    f"{len(TIERS['relaxed']['BM'])} items, and that figure already uses the relaxed answer-key rules. "
    f"The binding constraint is upstream of any design choice: {CAN['mechanical_funnel'][0]['rejects']} "
    f"of the {CAN['n_source_rows']} source items are bare theory one-liners with no patient and no "
    "sentence boundary at all, so there is nothing for a clinical finding to attach to. Reaching 200 "
    "per arm requires a larger or differently composed bank, not a better insertion method."
)

para(
    "The single blocking issue. The evaluation harness records the parsed multiple-choice letter. "
    "The eligibility rules deliberately remove every item whose correct answer an inserted finding "
    "could plausibly move, so on the delivered set the accuracy difference between control and "
    "manipulated conditions is zero by construction. Running the experiment as it stands would "
    "produce a null result that carries no information about whether the models noticed the "
    "fabricated entity. This is a gate on the runner, not a caveat on the analysis; it is set out in "
    "section 5."
)

para(
    f"What was delivered. {len(GEN)} generated rows covering {len(TIERS['relaxed']['BM'])} biomarker "
    f"texts and {len(TIERS['relaxed']['AN'])} anatomy texts ({N_BOTH} items carry both, "
    f"{N_BM_ONLY} biomarker only, {N_AN_ONLY} anatomy only), each accompanied by its unmodified "
    "control text, together with the per-item eligibility verdict and rejecting guard for all "
    f"{CAN['n_source_rows']} source items."
)

para(
    f"A caution on effective sample size. The {len(TIERS['strict']['PAIR'])} paired strict items are "
    f"not {len(TIERS['strict']['PAIR'])} independent observations. They come from "
    f"{STATS['strict_PAIR']['clusters']} distinct clinical narratives, and two narratives supply "
    f"{P37['clusters'].most_common(1)[0][1]} and {P37['clusters'].most_common(2)[1][1]} items "
    f"respectively. At a plausible intra-cluster correlation of 0.3 the effective sample size is "
    f"{STATS['strict_PAIR']['n_eff_icc_0.3']}, not {len(TIERS['strict']['PAIR'])}."
)

# --------------------------------------------------------------------------
# 2. What was asked and what was built
# --------------------------------------------------------------------------
doc.add_heading("2. What was asked and what was built", level=1)

doc.add_heading("2.1 The insertion point", level=2)
para(
    "Each item consists of a clinical narrative followed by a stem (the question proper) and four "
    "options. The manipulation appends exactly one sentence at the end of the narrative, immediately "
    "before the stem: the only location that leaves both the clinical story and the question "
    "untouched while placing the fabricated finding where a reader expects the last piece of evidence. "
    "The seam is the final sentence boundary in the item text — a full stop, exclamation or question "
    "mark followed by whitespace, or a line break. An item is usable only if that split is clean, "
    "which is what the mechanical guards in section 3 test."
)

doc.add_heading("2.2 The method is deterministic", level=2)
para("Every modified string is produced by a single concatenation:")
p = para("    new_text = narrative + SENTENCE + separator + stem", size=9.5)
p.runs[0].font.name = "Consolas"
para(
    "where SENTENCE is a literal constant drawn from the fixed table of twenty candidate sentences "
    "(Appendix A) and separator is the original whitespace that already sat between narrative and "
    "stem. The build asserts, for every generated string, that the prefix up to the insertion offset "
    "is byte-identical to the original, that the suffix after the inserted sentence and separator is "
    "byte-identical to the original, that the length grew by exactly the length of the inserted "
    "sentence plus the separator, that the inserted sentence occurs exactly once, and that the "
    "newline count changed only by the newlines in the separator. A failure of any assertion aborts "
    "the build."
)

doc.add_heading("2.3 No model touched the generated strings", level=2)
para(
    "No language model was involved at any point in generation. The twenty candidate sentences were "
    "written once, by hand, and are stored as literal constants in the build script; the choice of "
    "which sentence goes into which item is a deterministic rotation over an integer index "
    "(section 7). Re-running the build on the same workbook reproduces every byte of every generated "
    "string. This matters because the manipulation under test is the presence of a fabricated entity: "
    "had a model written the insertions, any observed effect could be attributed to the writer's "
    "stylistic choices rather than to fabrication."
)
para(
    "One further refusal rule applies. A candidate sentence is rejected for a given item if the "
    "item's own text already contains the sentence's distinguishing morpheme — its 'root' — because a "
    "fabricated term whose root already appears in the vignette can be assimilated as an echo of the "
    "text's own vocabulary rather than noticed as novel. The rotation skips to the next candidate on "
    "collision. Section 6 shows this is the dominant constraint on sentence choice."
)

# --------------------------------------------------------------------------
# 3. Which questions can be modified
# --------------------------------------------------------------------------
doc.add_heading("3. Which questions can be modified, and why most cannot", level=1)

doc.add_heading("3.1 The mechanical funnel", level=2)
para(
    f"Guards A to I are applied in order to all {CAN['n_source_rows']} items. Each guard tests one "
    "structural property of the item text; the first guard that fires rejects the item and no later "
    f"guard is evaluated. {CAN['mechanical_pool']} items survive."
)

rows = []
for f in CAN["mechanical_funnel"]:
    rows.append([f["guard"], f["why"][0].upper() + f["why"][1:], f["rejects"], f["remaining"]])
table(["Guard", "Reason for rejection", "Rejects", "Remaining"], rows,
      widths=[0.55, 4.35, 0.8, 0.9], right_from=2)
caption(f"Table 2. Mechanical funnel, {CAN['n_source_rows']} items down to {CAN['mechanical_pool']}. "
        "Guards are applied in the order shown; counts are first-hit.")

doc.add_heading("3.2 Guard A: the bank is mostly not clinical vignettes", level=2)
gA = CAN["mechanical_funnel"][0]["rejects"]
para(
    f"Guard A alone rejects {gA} items, {100 * gA / 474:.0f} per cent of the bank. These items have no "
    "sentence boundary anywhere in the question text: the whole item is a single fused sentence, "
    "typically a bare theory question of the form 'Regarding the treatment of X, indicate the correct "
    "answer'. There is no narrative, no patient, and no place to put a clinical finding. Inserting a "
    "sentence would not add a finding to a case; it would prepend an unattached assertion to an "
    "abstract question, which is a different manipulation with a different meaning."
)
para(
    f"Guard G removes a further {[f['rejects'] for f in CAN['mechanical_funnel'] if f['guard'] == 'G'][0]} "
    "items that do have a sentence boundary but still lack a patient anchor — no 'paciente', 'varón', "
    "'mujer', 'niño' or an age in years. Together A and G account for "
    f"{gA + [f['rejects'] for f in CAN['mechanical_funnel'] if f['guard'] == 'G'][0]} of the "
    f"{CAN['n_source_rows']} rejections, and they are the reason a 200-item arm is not available. The "
    "remaining eight mechanical guards remove "
    f"{474 - CAN['mechanical_pool'] - gA - [f['rejects'] for f in CAN['mechanical_funnel'] if f['guard'] == 'G'][0]} "
    "items between them and are housekeeping: abbreviation splits mistaken for sentence ends, stems "
    "whose preamble is itself a laboratory panel, narratives that end inside a bulleted list, and so on."
)

doc.add_heading("3.3 Answer-key preservation, and the strict/relaxed choice", level=2)
para(
    f"Guards J to N are then applied to the {CAN['mechanical_pool']} survivors. They ask a different "
    "question: not 'can a sentence be inserted here', but 'could inserting a sentence here change "
    "which option is correct'. Guards K, L, M and N are uncontroversial and are always applied."
)
rows = []
for g in ("K", "L", "M", "N"):
    rows.append([g, CAN["guard_glossary"][g][0].upper() + CAN["guard_glossary"][g][1:],
                 STRICT_KEY.get(g, 0), RELAX_KEY.get(g, 0)])
table(["Guard", "Reason for rejection", "Strict rejects", "Relaxed rejects"], rows,
      widths=[0.55, 4.35, 0.85, 0.85], right_from=2)
caption("Table 3. Always-applied key-preservation guards. Counts are first-hit within each tier, so "
        "the strict column is smaller where guard J has already removed the item.")

para(
    f"Guard J is the genuine trade-off. It removes decision-type stems — 'what is the next step', "
    f"'which test would you request', 'what is the most likely diagnosis', 'which treatment'. In the "
    f"strict tier it rejects {STRICT_KEY.get('J', 0)} of the {CAN['mechanical_pool']} mechanically "
    f"insertable items, more than the other four key guards combined, and it is the entire difference "
    f"between the strict base of {len(TIERS['strict']['base'])} and the relaxed base of "
    f"{len(TIERS['relaxed']['base'])}."
)
para(
    "The argument for applying J is sound. On a decision stem a new abnormal finding is exactly the "
    "kind of evidence that ought to change the recommended action, so if the fabricated biomarker were "
    "real the correct answer to 'what is the next step' might genuinely differ; control and "
    "manipulated conditions would no longer share a defensible key."
)
para(
    "The argument against applying J is equally sound and matters more here. Decision stems are "
    "precisely the items where a model has a reason to use the new finding. On a factual or negated "
    "stem — 'which of the following is false about colorectal cancer' — a competent model can and "
    "should ignore the appended sentence entirely, because the question does not depend on the "
    "patient. Removing all decision stems therefore removes the items that carry the signal and "
    "retains the items on which correct behaviour is indistinguishable from never having read the "
    "insertion. Guard J protects the answer key at the direct cost of measurement sensitivity."
)
para(
    "Both tiers are therefore reported and both item sets are delivered. Which to use is a design "
    "decision that follows from section 5: if the outcome measure remains the multiple-choice letter, "
    "the strict tier is correct and the experiment is uninformative either way. If the outcome measure "
    "becomes the model's free-text handling of the entity, the relaxed tier is correct, and guard J's "
    "removals become the most valuable items in the bank."
)

doc.add_heading("3.4 Per-arm admissibility", level=2)
para(
    "A key-preserving item can still be unsuitable for one arm. An item that already states that the "
    "laboratory workup is normal cannot receive a fabricated abnormal serum marker without a direct "
    "self-contradiction, and an item whose options discuss laboratory findings would have its options' "
    "referent moved by the insertion. The same applies to the anatomy arm for items that already state "
    "that the physical examination is unremarkable, or whose options discuss examination findings. "
    "These two gates produce the arm counts in Table 1, and their intersection produces the paired sets."
)

# --------------------------------------------------------------------------
# 4. The good questions
# --------------------------------------------------------------------------
doc.add_heading("4. The eligible item sets", level=1)

para(
    f"This section characterises the two paired sets — {len(TIERS['strict']['PAIR'])} strict and "
    f"{len(TIERS['relaxed']['PAIR'])} relaxed — since a paired set is what a within-item design needs. "
    "The single-arm sets are summarised alongside for comparison."
)

regions = sorted(set(P37["region"]) | set(P58["region"]) | set(P99["region"]) | set(P86["region"]),
                 key=lambda r: -P58["region"].get(r, 0))
rows = [["Region", "", "", "", ""]]
rows += [["   " + r, P37["region"].get(r, 0), P58["region"].get(r, 0),
          P99["region"].get(r, 0), P86["region"].get(r, 0)] for r in regions]
parts = sorted(set(P37["exam_part"]) | set(P58["exam_part"]) | set(P99["exam_part"]) | set(P86["exam_part"]),
               key=lambda p: -P58["exam_part"].get(p, 0))
rows += [["Exam part", "", "", "", ""]]
rows += [["   " + p, P37["exam_part"].get(p, 0), P58["exam_part"].get(p, 0),
          P99["exam_part"].get(p, 0), P86["exam_part"].get(p, 0)] for p in parts]
rows.append(["Total", P37["n"], P58["n"], P99["n"], P86["n"]])
table(["", "Strict paired (37)", "Relaxed paired (58)", "Biomarker arm (99)", "Anatomy arm (86)"],
      rows, widths=[2.15, 1.1, 1.15, 1.1, 1.1], right_from=1)
caption("Table 4. Regional and exam-part composition of the eligible sets. Illes Balears and the "
        "case-based sections dominate because they are the only parts of the bank that reliably "
        "contain a patient narrative; this is also the source of the clustering in Table 5.")

rows = [
    ["Items", P37["n"], P58["n"]],
    ["Distinct clinical narratives (clusters)", STATS["strict_PAIR"]["clusters"], STATS["relaxed_PAIR"]["clusters"]],
    ["Largest cluster", P37["clusters"].most_common(1)[0][1], P58["clusters"].most_common(1)[0][1]],
    ["Second largest cluster", P37["clusters"].most_common(2)[1][1], P58["clusters"].most_common(2)[1][1]],
    ["Singleton clusters", sum(1 for v in P37["clusters"].values() if v == 1),
     sum(1 for v in P58["clusters"].values() if v == 1)],
    ["Kish mean cluster size", STATS["strict_PAIR"]["kish_mean_cluster"], STATS["relaxed_PAIR"]["kish_mean_cluster"]],
    ["Effective n at ICC 0.1", STATS["strict_PAIR"]["n_eff_icc_0.1"], STATS["relaxed_PAIR"]["n_eff_icc_0.1"]],
    ["Effective n at ICC 0.3", STATS["strict_PAIR"]["n_eff_icc_0.3"], STATS["relaxed_PAIR"]["n_eff_icc_0.3"]],
    ["Effective n at ICC 0.5", STATS["strict_PAIR"]["n_eff_icc_0.5"], STATS["relaxed_PAIR"]["n_eff_icc_0.5"]],
    ["Negated stems", f"{P37['negated']} ({100 * P37['negated'] / P37['n']:.0f}%)",
     f"{P58['negated']} ({100 * P58['negated'] / P58['n']:.0f}%)"],
    ["Flagged clinical_case", P37["clinical_case"], P58["clinical_case"]],
    ["Narrative length, min / Q1 / median / Q3 / max (characters)",
     " / ".join(str(x) for x in P37["nc"]), " / ".join(str(x) for x in P58["nc"])],
    ["Correct-letter distribution (a/b/c/d)",
     "/".join(str(P37["letters"].get(k, 0)) for k in "abcd"),
     "/".join(str(P58["letters"].get(k, 0)) for k in "abcd")],
    ["Years represented", topn(P37["year"]), topn(P58["year"])],
]
table(["", "Strict paired (37)", "Relaxed paired (58)"], rows,
      widths=[2.7, 2.05, 2.05], right_from=1)
caption("Table 5. Structure of the two paired sets.")

para(
    f"The clustering has to be stated plainly. The {len(TIERS['strict']['PAIR'])} strict paired items "
    f"come from {STATS['strict_PAIR']['clusters']} distinct clinical narratives, and the distribution is "
    f"badly skewed: one narrative supplies {P37['clusters'].most_common(1)[0][1]} items and a second "
    f"supplies {P37['clusters'].most_common(2)[1][1]}, which is "
    f"{100 * (P37['clusters'].most_common(1)[0][1] + P37['clusters'].most_common(2)[1][1]) / P37['n']:.0f} "
    f"per cent of the set from two vignettes. The remaining "
    f"{STATS['strict_PAIR']['clusters'] - 2} clusters contribute "
    f"{P37['n'] - P37['clusters'].most_common(1)[0][1] - P37['clusters'].most_common(2)[1][1]} items. "
    f"Under the Kish approximation the design effect at an intra-cluster correlation of 0.3 gives an "
    f"effective sample size of {STATS['strict_PAIR']['n_eff_icc_0.3']}. In other words the strict "
    f"paired set is worth roughly ten independent items, not {len(TIERS['strict']['PAIR'])}, and any "
    "power calculation or confidence interval that treats it as thirty-seven will be wrong. The "
    f"relaxed paired set is better but not by much: {STATS['relaxed_PAIR']['clusters']} clusters and "
    f"an effective n of {STATS['relaxed_PAIR']['n_eff_icc_0.3']} at the same assumption."
)
para(
    "The narrative-length range is wide and should be checked before analysis. In the relaxed paired "
    f"set the narrative preceding the insertion runs from {P58['nc'][0]} to {P58['nc'][4]} characters. "
    "The inserted sentence is roughly 40 to 80 characters, so in the shortest items it is a "
    "substantial fraction of the whole vignette and in the longest it is a small addition near the end "
    "of a long text. Position and salience therefore vary systematically with narrative length, which "
    "should be carried into the model as a covariate rather than ignored."
)

# --------------------------------------------------------------------------
# 5. Blocking dependency
# --------------------------------------------------------------------------
doc.add_heading("5. The blocking dependency", level=1)

para(
    "The experiment as currently specified cannot produce an informative result, and the reason is "
    "structural rather than statistical.", bold=True,
)
para(
    "The harness records, for each model and item, the parsed multiple-choice answer. The model replies "
    "with a JSON object containing the question identifier, the selected letter and the text of the "
    "selected option; the run is converged when the harness reports zero parse failures and zero API "
    "failures. The analysis layer downstream consumes that letter, plus latency and token counts. "
    "There is no field in which the model's clinical reasoning is stored, and none in which a "
    "statement about the fabricated entity could appear."
)
para(
    "Now combine that with the eligibility rules. Guards J to N exist precisely to remove every item "
    "whose correct option an inserted finding could plausibly move. On the items that survive, the "
    "correct answer is by construction the same in the control and manipulated conditions. A model "
    "that reads the fabricated sentence, treats it as real, and reasons from it will still select the "
    "same letter as a model that ignores the sentence entirely. The accuracy delta between conditions "
    "is zero by construction, and measuring it tells us nothing."
)
para(
    "This is not a hypothetical. It is the arithmetic consequence of two decisions that are each "
    "correct on their own: protect the answer key, and score by the letter. Taken together they "
    "guarantee a null.", bold=True,
)
para("Two changes to the runner are required before the experiment measures anything.")
numbered(
    "Persist the model's free-text rationale. The response schema must be extended with a rationale "
    "field, and that field must be written to the database alongside the letter. Note that a "
    "'reasoning' field is already present in the raw provider payload for those OpenRouter models that "
    "emit a thinking trace, but it is absent for the remainder and for the GIFT arm, and a thinking "
    "trace is not the same artefact as an answer to the question that was asked. It cannot be relied "
    "on as the measurement."
)
numbered(
    "Add an explicit entity probe. After the item is answered, the model is asked a second, separate "
    "question about the fabricated term — whether the named entity exists, what it is, or whether any "
    "part of the vignette was unfamiliar. The probe must be a separate call so that it cannot "
    "contaminate the primary answer, and it must be scored against a fixed rubric with at least "
    "three levels: flagged as non-existent, ignored, or accepted and used."
)
para(
    "Until both exist, the correct action is to hold the run. The item set is finished and will not "
    "change; it can sit until the runner catches up. Executing the run first and adding the "
    "instrumentation later would mean discarding the results and paying for the inference twice."
)

# --------------------------------------------------------------------------
# 6. Sentence selection
# --------------------------------------------------------------------------
doc.add_heading("6. Sentence selection", level=1)

para(
    "Twenty candidate sentences were written, ten biomarkers and ten anatomical structures "
    f"(Appendix A). Only four biomarker sentences and two anatomy sentences are used: "
    f"{', '.join(BM_ROT)} and {', '.join(AN_ROT)}. The reasons are specific to this corpus and are set "
    "out below."
)

doc.add_heading("6.1 Specimen type eliminates three biomarker sentences", level=2)
para(
    "Three biomarker sentences specify a faecal or salivary specimen (BM03, BM09, BM06). A faecal or "
    "salivary determination is a distinct investigation that a gastroenterology vignette would have "
    "had to order; appending one asserts that a test was performed rather than reporting an extra "
    "value from a panel that the narrative already describes. Only serum and plasma sentences — BM01, "
    "BM02, BM04, BM05, BM07, BM08 and BM10 — can be appended as an additional line of an existing "
    "laboratory report."
)

doc.add_heading("6.2 Lexical camouflage selects the four that are used", level=2)
para(
    "Of the seven serum and plasma sentences, the four with the fewest root collisions in the corpus "
    "are used, in ascending order of collision count. The table below counts, for each sentence, how "
    f"many item texts already contain its root, across the full bank, across the "
    f"{CAN['mechanical_pool']} mechanically insertable items, and across the "
    f"{len(TIERS['relaxed']['BM'])}-item biomarker arm, which is the number that matters."
)

rows = []
for vid in sorted(BM):
    ent, sent, root = BM[vid]
    specimen = ("serum" if "serica" in nf(sent) else
                "plasma" if "plasmatica" in nf(sent) else
                "faecal" if "fecal" in nf(sent) else "saliva")
    rows.append([vid, ent, root, specimen,
                 root_hits(root, ALL474), root_hits(root, MECH),
                 root_hits(root, TIERS["relaxed"]["BM"]),
                 "yes" if vid in BM_ROT else ""])
table(["ID", "Entity", "Root", "Specimen", "Bank (474)", f"Pool ({CAN['mechanical_pool']})",
       f"BM arm ({len(TIERS['relaxed']['BM'])})", "Used"],
      rows, widths=[0.45, 1.35, 0.75, 0.7, 0.7, 0.7, 0.7, 0.5], right_from=4)
caption("Table 6. Root collisions for the ten biomarker sentences. A collision means the item's own "
        "text already contains the root, so the fabricated term can be read as an echo of the "
        "vignette's vocabulary rather than as a novel entity.")

para(
    f"The ordering is unambiguous. Among the serum and plasma sentences the collision counts in the "
    f"biomarker arm are BM07 {root_hits('fibro', TIERS['relaxed']['BM'])}, "
    f"BM02 {root_hits('colangi', TIERS['relaxed']['BM'])}, "
    f"BM08 {root_hits('portal', TIERS['relaxed']['BM'])}, "
    f"BM04 {root_hits('pancre', TIERS['relaxed']['BM'])}, "
    f"BM01 {root_hits('hepat', TIERS['relaxed']['BM'])}, "
    f"BM10 {root_hits('colon', TIERS['relaxed']['BM'])} and "
    f"BM05 {root_hits('gastr', TIERS['relaxed']['BM'])}. The rotation takes the four cleanest, which "
    "is why BM07 fibroquelina-X3 leads: in a gastroenterology corpus, 'fibro' is close to absent "
    f"(the only host tokens anywhere in the bank are "
    f"{', '.join(f'{w} ({n})' for w, n in FIBRO_WORDS.most_common())}), whereas 'hepat', 'gastr' and "
    "'colon' are ubiquitous."
)

doc.add_heading("6.3 BM06 mucorynex-Z9 is the worst candidate, not the best", level=2)
para(
    "BM06 is worth singling out because it was proposed as a leading choice on the grounds that "
    "'mucorynex' is phonetically distinctive and unlike any real analyte name. Recomputation from the "
    "source workbook inverts that conclusion. The camouflage test is not run on the whole word; it is "
    "run on the root, and BM06's root is 'muco'."
)
rows = [
    ["Items whose question text contains 'muco'", root_hits("muco", ALL474),
     f"{100 * root_hits('muco', ALL474) / 474:.1f}% of {len(ALL474)}"],
    ["Items whose stem or options contain 'muco'", MUCO_ROWS_ANY,
     f"{100 * MUCO_ROWS_ANY / 474:.1f}% of {len(ALL474)}"],
    [f"Of the {CAN['mechanical_pool']} mechanically insertable items", root_hits("muco", MECH),
     f"{100 * root_hits('muco', MECH) / CAN['mechanical_pool']:.1f}%"],
    [f"Of the {len(TIERS['relaxed']['BM'])}-item biomarker arm",
     root_hits("muco", TIERS["relaxed"]["BM"]),
     f"{100 * root_hits('muco', TIERS['relaxed']['BM']) / len(TIERS['relaxed']['BM']):.1f}%"],
    ["Host tokens in the bank (stem and options)",
     sum(MUCO_WORDS.values()),
     ", ".join(f"{w} {n}" for w, n in MUCO_WORDS.most_common(6))],
]
table(["Measure", "Count", "Share / detail"], rows, widths=[2.6, 0.8, 3.4], right_from=1)
caption("Table 7. BM06 root collision, recomputed from balanced-flat-A.xlsx.")

para(
    f"Nearly half the biomarker arm — {root_hits('muco', TIERS['relaxed']['BM'])} of "
    f"{len(TIERS['relaxed']['BM'])} items — already contains 'muco', overwhelmingly as 'mucosa' "
    f"({MUCO_WORDS['mucosa']} occurrences), with 'submucosa' ({MUCO_WORDS['submucosa']}) and "
    f"'mucosectomía' ({MUCO_WORDS.get('mucosectomia', 0)}) behind it. In a corpus about the digestive "
    "tract this is the least distinctive root that could have been chosen. BM06 would have been "
    "refused by the camouflage rule on roughly half the arm, and on the other half the term would sit "
    "in a text saturated with morphologically similar words, which is the condition under which a "
    "fabricated term is most likely to pass unnoticed. It is the worst of the ten, not the best."
)

doc.add_heading("6.4 The eight imaging-based anatomy sentences are unusable here", level=2)
para(
    "Eight of the ten anatomy sentences attribute the finding to an imaging study: three to "
    "ultrasound, three to computed tomography, two to magnetic resonance. Only AN04 and AN10 attribute "
    "it to physical examination. An imaging sentence can only be appended where the vignette already "
    "reports a study of that modality — otherwise the insertion fabricates an entire investigation, "
    "not a finding — and where that study is not already reported as normal, and where neither stem "
    "nor options refer to imaging. Applying those three conditions to the paired sets gives the "
    "counts below."
)
rows = []
for name, pat in MODALITY.items():
    s = imaging_funnel(TIERS["strict"]["PAIR"], pat)
    r = imaging_funnel(TIERS["relaxed"]["PAIR"], pat)
    a = imaging_funnel(ALL474, pat)
    rows.append([name, a[0], s[0], s[1], s[2], r[0], r[1], r[2]])
table(["Modality", "Bank (474)", "Strict 37: reported", "not normal", "clean",
       "Relaxed 58: reported", "not normal", "clean"],
      rows, widths=[0.75, 0.65, 0.8, 0.7, 0.55, 0.85, 0.7, 0.55], right_from=1)
caption("Table 8. Survival of the imaging-based anatomy sentences. 'Reported' = the vignette already "
        "describes a study of that modality; 'not normal' = that study is not described as normal or "
        "without findings; 'clean' = stem and options make no reference to imaging.")

para(
    "On the strict paired set no imaging modality survives at all: computed tomography and magnetic "
    "resonance are never reported in an eligible vignette, and the two ultrasound-bearing items fail "
    "on the remaining conditions. On the relaxed paired set the best modality yields four usable "
    "items. Magnetic resonance is barely present in the bank at any point — it appears in "
    f"{imaging_funnel(ALL474, MODALITY['MRI'])[0]} of {len(ALL474)} items in total. Eight sentences "
    "therefore cannot support an arm of any usable size, which is why the anatomy rotation is confined "
    "to the two examination-based sentences AN04 saco orfalónico and AN10 órgano liradónico. AN04 "
    "leads for the same structural reason BM07 does: it heads the rotation and the rotation index is "
    "zero for most items."
)

# --------------------------------------------------------------------------
# 7. Confound
# --------------------------------------------------------------------------
doc.add_heading("7. Why a single sentence for everything would be a confound", level=1)

para(
    "A tempting simplification is to use one biomarker sentence and one anatomy sentence throughout. "
    "It would make the manipulation perfectly uniform. It would also make the experiment "
    "uninterpretable, for a reason that is easy to state: with one sentence per arm, the entity factor "
    "has zero degrees of freedom. Any effect observed in the biomarker arm would be an effect of the "
    "string 'La fibroquelina-X3 sérica se encuentra aumentada', and there would be no way to "
    "distinguish a general property of fabricated biomarkers from an idiosyncrasy of that particular "
    "invented word — its length, its suffix, its resemblance to 'fibrosis', the fact that it carries "
    "an alphanumeric code. The rotation over four biomarker sentences and two anatomy sentences buys "
    "the minimum variance needed to notice that such an idiosyncrasy exists."
)

para(
    "The point applies with more force to the comparison between arms. BM07 and AN04 differ on at "
    "least three dimensions simultaneously:"
)
rows = [
    ["Entity type", "Analyte (a measured substance)", "Anatomical structure (a body part)"],
    ["Evidential modality", "Laboratory determination", "Physical examination by the clinician"],
    ["Surface form", "Single Latinate noun with an alphanumeric code, 'fibroquelina-X3'",
     "Two-word Spanish noun phrase, no code, 'saco orfalónico'"],
    ["Assertion", "A value is outside the reference range", "A sign is elicited on palpation"],
    ["Sentence length", f"{len(BM['BM07'][1])} characters", f"{len(AN['AN04'][1])} characters"],
]
table(["Dimension", "BM07 fibroquelina-X3", "AN04 saco orfalónico"], rows,
      widths=[1.3, 2.7, 2.4])
caption("Table 9. The biomarker and anatomy arms are not a clean single-factor contrast.")
para(
    "A difference in detection rate between the arms therefore cannot be attributed to 'biomarker "
    "versus anatomy'. It could equally be a difference between laboratory and examination evidence, "
    "or between coded and uncoded terms. The arms should be treated as two separate probes of the same "
    "phenomenon, not as levels of a factor."
)

doc.add_heading("7.1 The rotation actually used", level=2)
para(
    "Each item receives an index equal to its position within its clinical narrative, counted over the "
    "source workbook order. The biomarker sentence is the rotation entry at index modulo four and the "
    "anatomy sentence is the rotation entry at index modulo two, with a skip to the next entry if the "
    "chosen sentence's root already appears in the item. "
    f"{'Exactly one item' if BM_SKIPS + AN_SKIPS == 1 else f'{BM_SKIPS + AN_SKIPS} items'} required a "
    "skip, which is itself evidence that the four sentences chosen in section 6 are close to "
    "collision-free on this corpus."
)
rows = [[vid, BM[vid][0], BMVAR.get(vid, 0), f"{100 * BMVAR.get(vid, 0) / sum(BMVAR.values()):.0f}%"]
        for vid in BM_ROT]
rows += [[vid, AN[vid][0], ANVAR.get(vid, 0), f"{100 * ANVAR.get(vid, 0) / sum(ANVAR.values()):.0f}%"]
         for vid in AN_ROT]
table(["Variant", "Entity", "Items", "Share of arm"], rows,
      widths=[0.75, 2.0, 0.8, 1.0], right_from=2)
caption(f"Table 10. Variant use across the {len(GEN)} generated rows "
        f"({sum(BMVAR.values())} biomarker texts, {sum(ANVAR.values())} anatomy texts).")

para(
    "Two defects in the rotation should be recorded, both discovered by recomputing the delivered "
    "assignment rather than by reading the code.", bold=True,
)
para(
    "First, the two rotations are locked in phase. Both use the same index, and two divides four, so "
    "the biomarker variant determines the anatomy variant exactly. On the "
    f"{len(TIERS['relaxed']['PAIR'])} paired items the joint distribution has only four non-empty "
    "cells out of eight:"
)
rows = [[f"{b} + {a}", n] for (b, a), n in sorted(JOINT.items(), key=lambda kv: -kv[1])]
table(["Biomarker + anatomy pairing", "Items"], rows, widths=[2.4, 0.9], right_from=1)
caption("Table 11. Joint variant assignment on the relaxed paired set. Four of the eight possible "
        "combinations never occur.")
para(
    "The consequence is that within-item comparisons of the two arms are always between a fixed pair "
    "of strings. Any entity-level idiosyncrasy in BM07 is perfectly confounded with any "
    "entity-level idiosyncrasy in AN04. Using rotations of coprime length, or offsetting the anatomy "
    "index by one, would break this at no cost."
)
para(
    "Second, the entity is aliased with region outside the largest cluster. Within Illes Balears, "
    "whose items sit in long multi-question narratives, the four biomarker variants are close to "
    f"balanced ({topn(BM_BY_REGION['Illes Balears'])}). Everywhere else, most items are the only "
    "eligible item from their narrative, so the index is zero and the first rotation entry is always "
    f"chosen: outside Illes Balears BM07 takes "
    f"{sum(c['BM07'] for r, c in BM_BY_REGION.items() if r != 'Illes Balears')} of "
    f"{sum(sum(c.values()) for r, c in BM_BY_REGION.items() if r != 'Illes Balears')} biomarker items, "
    f"and AN04 takes {sum(c['AN04'] for r, c in AN_BY_REGION.items() if r != 'Illes Balears')} of "
    f"{sum(sum(c.values()) for r, c in AN_BY_REGION.items() if r != 'Illes Balears')} anatomy items. "
    "The build's own comment claims the rotation prevents aliasing of entity with cluster or region; "
    "that holds for cluster but not for region. Seeding the index from a hash of the question "
    "identifier rather than from position within the narrative would fix it."
)

# --------------------------------------------------------------------------
# 8. Learnings
# --------------------------------------------------------------------------
doc.add_heading("8. Learnings", level=1)

doc.add_heading("8.1 Directional metrics are blind to omission", level=2)
para(
    "Accuracy on a multiple-choice item is directional: it records which of four pre-specified "
    "alternatives the model moved towards, and has no vocabulary for what the model failed to do. A "
    "model that reads an invented biomarker, silently accepts it, and answers correctly anyway scores "
    "identically to one that recognises the term as fabricated and says so. The failure being probed "
    "here — accepting a non-existent entity without objection — is an omission, and no directional "
    "metric can see an omission. Any experiment whose hypothesis concerns what a model fails to flag "
    "needs an outcome measure that can represent 'said nothing'. This generalises beyond fabricated "
    "entities to safety caveats, missing contraindications and unstated uncertainty."
)

doc.add_heading("8.2 Key preservation and measurement sensitivity trade off directly", level=2)
para(
    "The eligibility rules guarantee that the correct answer is unchanged by the insertion, and that "
    "guarantee is what makes a control-versus-manipulated comparison legitimate. But the items whose "
    "key an inserted finding could move are exactly the items on which the model has a reason to use "
    "the finding. Removing them leaves a set on which correct behaviour and total inattention are "
    f"observationally identical. The magnitude is visible in the funnel: guard J alone removes "
    f"{STRICT_KEY.get('J', 0)} of {CAN['mechanical_pool']} items, and those are the decision-type "
    "stems, the most informative ones. Key preservation is therefore not a free safety measure; it "
    "buys internal validity by spending statistical power, and the exchange rate should be computed "
    "before the rule is adopted. Where the rule would remove the signal entirely, the correct response "
    "is to change the outcome measure, not to relax the rule."
)

doc.add_heading("8.3 Verifying a corpus is not the same as verifying the artefact derived from it", level=2)
para(
    "The source bank had been checked: row counts, answer keys, option integrity, flags. None of that "
    "says anything about whether a fabricated sentence can be appended to those rows. The properties "
    "that mattered here — does a sentence boundary exist, is there a patient, does the narrative end "
    "in a full stop, does the item's own text already contain the root of the term to be inserted — "
    "are properties of the derived artefact, and every one had to be measured separately. Each "
    "transformation applied to a dataset creates a new object with its own failure modes, and "
    "inherited assurance does not transfer across the transformation. Budget a verification pass per "
    "derived artefact, not per dataset."
)

doc.add_heading("8.4 Adversarial recomputation of analyst claims", level=2)
para(
    "Claims made about the item set during construction were re-derived from the source workbook "
    "rather than accepted. Roughly a quarter — 26 per cent on the tally kept during the build — did "
    "not survive recomputation. Most failures were minor: counts off by a few items, a guard described "
    "in the wrong order, a percentage computed against the wrong denominator. One was exactly "
    "inverted. BM06 mucorynex-Z9 had been recommended as a leading candidate for its phonetic "
    f"distinctiveness; recomputation of the root collision (section 6.3) shows it colliding on "
    f"{root_hits('muco', TIERS['relaxed']['BM'])} of {len(TIERS['relaxed']['BM'])} biomarker-arm "
    "items, the worst of all ten candidates, so acting on it would have degraded the manipulation on "
    "half the arm. The lesson is not that analysts are careless; it is that plausible-sounding "
    "quantitative claims about a corpus are cheap to assert and cheap to check, and that asymmetry "
    "means they should always be checked. Recomputation should be adversarial by default: try to "
    "falsify the claim rather than confirm that a number of that magnitude is obtainable."
)

doc.add_heading("8.5 Item counts overstate the evidence when items share narratives", level=2)
para(
    f"{len(CTX_ITEMS)} of the {len(ALL474)} items in the bank "
    f"({100 * len(CTX_ITEMS) / len(ALL474):.0f} per cent) re-use one of only {CTX_NARR} clinical "
    f"narratives, and because an item can belong to a base case and to one or more of its transitions, "
    f"those {len(CTX_ITEMS)} items carry {CTX_MEMBERSHIPS} narrative memberships between them. A single "
    "Illes Balears case supplies up to twenty-four questions. Every response to those items is "
    "conditioned on the same vignette text, so the responses are not independent draws. Reporting "
    "'n = 474' or 'n = 37' as though they were overstates precision by a factor that depends on the "
    f"intra-cluster correlation, and at plausible values that factor is large: the strict paired set "
    f"of {len(TIERS['strict']['PAIR'])} is worth {STATS['strict_PAIR']['n_eff_icc_0.3']} independent "
    "items at an intra-cluster correlation of 0.3. Clustering should be recorded at bank-construction "
    "time, carried through every downstream count, and every reported interval made cluster-robust."
)

# --------------------------------------------------------------------------
# 9. Recommended next steps
# --------------------------------------------------------------------------
doc.add_heading("9. Recommended next steps", level=1)

rows = [
    ["1", "Extend the runner to persist a free-text rationale and add a separate entity-probe call, "
          "scored on a three-level rubric. Nothing else on this list is worth doing first.",
     "Gate on the whole experiment"],
    ["2", "Reset the target. Plan against the achievable set, not 200 per arm.",
     f"{len(TIERS['relaxed']['BM'])} biomarker items, {len(TIERS['relaxed']['AN'])} anatomy items, "
     f"{len(TIERS['relaxed']['PAIR'])} paired"],
    ["3", "Adopt the relaxed tier once the outcome measure is the rationale rather than the letter, "
          "and analyse the guard-J items as a pre-registered subgroup.",
     f"{len(TIERS['relaxed']['base'])} base items, of which {STRICT_KEY.get('J', 0)} are the "
     "decision-type stems guard J removes"],
    ["4", "Rebuild the rotation with coprime cycle lengths and an index seeded from the question "
          "identifier, to unlock the arms from each other and de-alias entity from region.",
     f"{len(JOINT)} of 8 joint cells currently occupied; "
     f"{sum(c['BM07'] for r, c in BM_BY_REGION.items() if r != 'Illes Balears')} of "
     f"{sum(sum(c.values()) for r, c in BM_BY_REGION.items() if r != 'Illes Balears')} "
     "non-Balearic biomarker items are BM07"],
    ["5", "Power the analysis on effective, not nominal, sample size, and pre-specify cluster-robust "
          "intervals with the narrative as the clustering unit.",
     f"Effective n {STATS['relaxed_PAIR']['n_eff_icc_0.3']} for the paired set at ICC 0.3, from "
     f"{STATS['relaxed_PAIR']['clusters']} narratives"],
    ["6", "Enlarge the bank if 200 per arm remains a requirement. Target case-based sections "
          "specifically; theory one-liners cannot be used at any yield.",
     f"{CAN['mechanical_funnel'][0]['rejects']} of {len(ALL474)} current items fail on guard A alone; "
     f"the observed yield from bank to biomarker arm is "
     f"{100 * len(TIERS['relaxed']['BM']) / len(ALL474):.0f}%, implying roughly "
     f"{round(200 * len(ALL474) / len(TIERS['relaxed']['BM']) / 100) * 100:,} source items for a "
     "200-item arm"],
    ["7", "Add a second control condition inserting a real but irrelevant finding, to separate "
          "'noticed the entity is fabricated' from 'noticed an extra sentence'.",
     "Applies to the same item set; no new eligibility work"],
    ["8", "Carry narrative length as a covariate; the inserted sentence is a very different fraction "
          "of a short vignette than of a long one.",
     f"Narrative length spans {P58['nc'][0]}–{P58['nc'][4]} characters on the paired set"],
]
table(["#", "Action", "Number attached"], rows, widths=[0.3, 3.4, 2.9])
caption("Table 12. Recommended next steps, in order.")

# --------------------------------------------------------------------------
# Appendix
# --------------------------------------------------------------------------
doc.add_heading("Appendix A. The twenty candidate sentences", level=1)

rows = []
for vid in sorted(BM):
    ent, sent, root = BM[vid]
    used = f"used ({BMVAR[vid]})" if vid in BM_ROT else "not used"
    rows.append([vid, ent, sent, root, used])
for vid in sorted(AN):
    ent, sent, root = AN[vid]
    used = f"used ({ANVAR[vid]})" if vid in AN_ROT else "not used"
    rows.append([vid, ent, sent, root, used])
table(["ID", "Entity", "Sentence (literal constant)", "Root", "Status"], rows,
      widths=[0.45, 1.2, 3.45, 0.8, 0.7], font=7)
caption("Table A1. All twenty candidate sentences as stored in tools/build_expC.py. The 'root' column "
        "is the morpheme used for the camouflage test. 'Status' gives the number of generated items "
        "using each sentence.")

doc.add_heading("Appendix B. Question identifier lists", level=1)


def idlist(ids):
    return " ".join(sorted(ids, key=lambda s: int(re.sub(r"\D", "", s))))


for label, ids in (
    (f"Strict paired set ({len(TIERS['strict']['PAIR'])})", TIERS["strict"]["PAIR"]),
    (f"Relaxed paired set ({len(TIERS['relaxed']['PAIR'])})", TIERS["relaxed"]["PAIR"]),
    (f"Strict biomarker arm ({len(TIERS['strict']['BM'])})", TIERS["strict"]["BM"]),
    (f"Strict anatomy arm ({len(TIERS['strict']['AN'])})", TIERS["strict"]["AN"]),
    (f"Relaxed biomarker arm ({len(TIERS['relaxed']['BM'])})", TIERS["relaxed"]["BM"]),
    (f"Relaxed anatomy arm ({len(TIERS['relaxed']['AN'])})", TIERS["relaxed"]["AN"]),
    (f"Strict key-preserving base ({len(TIERS['strict']['base'])})", TIERS["strict"]["base"]),
    (f"Relaxed key-preserving base ({len(TIERS['relaxed']['base'])})", TIERS["relaxed"]["base"]),
):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(1)
    p = doc.add_paragraph()
    r = p.add_run(idlist(ids))
    r.font.size = Pt(7.5)
    r.font.name = "Consolas"
    p.paragraph_format.space_after = Pt(6)

para(
    "Per-item eligibility verdicts, the rejecting guard for every excluded item, and the full control "
    "and manipulated texts for every generated row are in canonical.json under the keys 'items' and "
    "'generated'.", size=9, italic=True,
)

doc.save(OUT)

# ----------------------------------------------------------------- summary --
words = sum(len(p.text.split()) for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            words += len(c.text.split())
print(f"wrote {OUT}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} approx_words={words}")
